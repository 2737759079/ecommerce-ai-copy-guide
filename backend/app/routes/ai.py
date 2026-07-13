from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from app.database import get_db
from app.models.user import User
from app.models.product import Product
from app.models.live_script import LiveScript
from app.schemas import AICopyRequest, AIScriptRequest, AIResponse
from app.utils.security import require_merchant
from app.services import ai_service

router = APIRouter()


@router.post("/copy", response_model=AIResponse)
def generate_copy(payload: AICopyRequest, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    product = db.query(Product).filter(Product.id == payload.product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="商品不存在")
    data = {
        "id": product.id,
        "name": product.name,
        "category": product.category,
        "description": product.description,
        "price": product.price,
        "specs": product.specs or [],
    }
    result = ai_service.generate_product_copy(data, payload.style)
    product.ai_title = result.get("title", "")
    product.ai_selling_points = result.get("selling_points", "")
    product.ai_detail = result.get("detail", "")
    product.ai_slogan = result.get("slogan", "")
    db.commit()
    return AIResponse(result=result)


@router.post("/script", response_model=AIResponse)
def generate_script(payload: AIScriptRequest, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    product = db.query(Product).filter(Product.id == payload.product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="商品不存在")
    data = {
        "id": product.id,
        "name": product.name,
        "category": product.category,
        "description": product.description,
        "price": product.price,
        "specs": product.specs or [],
    }
    content = ai_service.generate_live_script(data, payload.style, payload.platform)
    script = LiveScript(product_id=product.id, title=f"{product.name}-{'直播' if payload.platform == 'live' else '短视频'}脚本", style=payload.style, content=content)
    db.add(script)
    db.commit()
    return AIResponse(result={"content": content, "title": f"{product.name}-{'直播' if payload.platform == 'live' else '短视频'}脚本"})


@router.post("/script/export")
def export_script(payload: dict, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    fmt = payload.get("format", "txt")
    content = payload.get("content", "")
    title = payload.get("title", "脚本")
    if fmt == "docx":
        from docx import Document
        import io
        doc = Document()
        doc.add_heading(title, level=1)
        for line in content.split("\n"):
            doc.add_paragraph(line)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        from fastapi.responses import StreamingResponse
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={title}.docx"})
    else:
        import io
        buffer = io.BytesIO(content.encode("utf-8"))
        from fastapi.responses import StreamingResponse
        return StreamingResponse(buffer, media_type="text/plain", headers={"Content-Disposition": f"attachment; filename={title}.txt"})


@router.get("/scripts")
def list_scripts(product_id: int = None, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    q = db.query(LiveScript)
    if product_id:
        q = q.filter(LiveScript.product_id == product_id)
    return q.order_by(LiveScript.created_at.desc()).all()
